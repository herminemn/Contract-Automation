from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.http import HttpResponse, Http404
from .forms import UploadFileForm, VariablesForm
from .functions import docx_words_replace
from .models import DocFile
from docx import Document
import re
from django.core.exceptions import ValidationError
import os
from django.core.files import File
from django.core.files.storage import default_storage
from contractautomation.settings import MEDIA_ROOT
from sendfile import sendfile
from xsendfile import XSendfileApplication


def files_list(request):
    uploaded_files = DocFile.objects.all()
    return render(request, 'uploads/files_list.html', {
        'uploaded_files': uploaded_files
    })


def upload_file(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('files_list')
    else:
        form = UploadFileForm()
    return render(request, 'uploads/upload.html', {
        'form': form
    })


def delete_file(request, pk):
    if request.method == 'POST':
        uploaded_file = DocFile.objects.get(pk=pk)
        uploaded_file.delete()
    return redirect('files_list')


# def new_files_list(request):
#     new_files = NewDocFile.objects.all()
#     return render(request, 'uploads/new_files_list.html', {
#         'new_files': new_files
#     })
#
#
# def delete_new_file(request, pk):
#     if request.method == 'POST':
#         new_file = NewDocFile.objects.get(pk=pk)
#         new_file.delete()
#     return redirect('new_files_list')


def edit_file(request, upload_id):
    instance = get_object_or_404(DocFile, id=upload_id)
    document = Document(instance.agreement)
    regex = re.compile(r"\{(.*?)\}")
    variables = []
    for paragraph in document.paragraphs:
        match = re.findall(regex, paragraph.text)
        variables.append(match)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                match = re.findall(regex, cell.text)
                variables.append(match)
    temp_list = []
    for variable in variables:
        for var in variable:
            if len(var) > 0:
                temp_list.append(var)
    variables = temp_list
    variables_set = sorted(set(temp_list), key=temp_list.index)
    inputs_list = []
    form_field = VariablesForm(request.POST, variables=variables_set)
    if request.method == 'POST':
        input_texts = form_field.get_input_text()
        print(form_field)
        for i, input_text in input_texts:
            inputs_list.append(input_text)
    # print(inputs_list)
    my_dict = dict(zip(variables_set, inputs_list))
    # print(my_dict)
    for word, replacement in my_dict.items():
        word_re = re.compile(word)
        docx_words_replace(document, word_re, replacement)
    regex1 = re.compile(r"{")
    regex2 = re.compile(r"}")
    replace = r""
    docx_words_replace(document, regex1, replace)
    docx_words_replace(document, regex2, replace)
    if len(inputs_list) != 0:
        contract_name = inputs_list[0]
        print(contract_name)
        file_name = '{}.docx'.format(contract_name)
        desktop = os.path.normpath(os.path.expanduser("~/Desktop"))
        document.save(os.path.join(desktop, file_name))
        messages.success(request, 'Saved on Desktop')
        # desktop_linux_mac = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop')
        # desktop_windows = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        # if os.name == 'nt':
        #     document.save(os.path.join(desktop_windows, file_name))
        # else:
        #     document.save(os.path.join(desktop_linux_mac, file_name))
        # response = HttpResponse(document, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        # response['Content-Disposition'] = 'attachment; filename="{}"'.format(file_name)
        # document.save(response)
        # return response
        # return redirect('new_files_list')
    return render(request, 'uploads/file_detail.html', {
        'variables': variables, 'form_field': form_field
    })


# def download_document(request, document):
#     file_path = os.path.join(MEDIA_ROOT, document)
#     if os.path.exists(file_path):
#         return sendfile(request, document.file.path)
#     return redirect('new_files_list')
#     # if os.path.exists(file_path):
    #     with open(file_path, 'rb') as fh:
    #         response = HttpResponse(fh.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    #         response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
    #         return response
    # raise Http404

#
# def new_files_list(request):
#     path = MEDIA_ROOT
#     new_files_list = []
#     for file in os.listdir(path):
#         if file.endswith('.doc') or file.endswith('.docx'):
#             new_files_list.append(os.path.basename(file))
#     return render(request, 'uploads/new_files_list.html', {
#         'new_files': new_files_list
#     })


# def delete_new_file(request, path, pk):
#     if request.method == 'POST':
#         file_path = os.path.join(MEDIA_ROOT, path)
#         if os.path.exists(file_path):
#
#             print(file_path)
#             os.remove(file_path)
#     return redirect('new_files_list')
